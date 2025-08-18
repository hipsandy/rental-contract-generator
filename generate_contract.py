import streamlit as st
import re
from docx import Document
from io import BytesIO
from collections import OrderedDict
from docx.shared import RGBColor

def extract_placeholders(docx_file):
    doc = Document(docx_file)
    placeholders = []
    for p in doc.paragraphs:
        for match in re.finditer(r"\{\{(.*?)\}\}", p.text):
            placeholders.append(match.group(1))
    # Deduplicate but preserve order
    seen = set()
    ordered_placeholders = []
    for ph in placeholders:
        if ph not in seen:
            ordered_placeholders.append(ph)
            seen.add(ph)
    return ordered_placeholders

def replace_in_paragraph(paragraph, values):
    """Replace placeholders in a paragraph (even if split across runs).
       Keeps original formatting for normal text, but replacements are bold + red.
    """
    if not values:
        return

    # Get the full paragraph text (merging runs)
    full_text = "".join(run.text for run in paragraph.runs)
    if not re.search(r"\{\{.*?\}\}", full_text):
        return  # no placeholders here

    # Split into text and placeholders
    parts = re.split(r"(\{\{.*?\}\})", full_text)

    # Clear all runs
    for run in paragraph.runs:
        run.text = ""
    paragraph._element.clear_content()

    # Rebuild runs
    for part in parts:
        if re.match(r"\{\{(.*?)\}\}", part):
            key = re.match(r"\{\{(.*?)\}\}", part).group(1)
            value = values.get(key, part)
            run = paragraph.add_run(value)
            # run.bold = True  Commenting for now since bold may take up more space.
            run.font.color.rgb = RGBColor(255, 0, 0)  # red
        else:
            if part:  # normal text
                paragraph.add_run(part)

def fill_template(docx_file, values):
    doc = Document(docx_file)

    for p in doc.paragraphs:
        replace_in_paragraph(p, values)

    for table in doc.tables:  # Also handle placeholders inside tables
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, values)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

st.title("Contract Generator")

uploaded_file = st.file_uploader("Upload Template with placeholders (.docx)", type="docx")
if uploaded_file:
    placeholders = extract_placeholders(uploaded_file)
    user_values = OrderedDict()
    for ph in placeholders:
        user_values[ph] = st.text_input(f"{ph}")

    if all(v not in ("", None) for v in user_values.values()):
        st.download_button(
            "Download Contract",
            data=fill_template(uploaded_file, user_values),
            file_name="contract_filled.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )